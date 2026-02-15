import os
import sys
import zipfile
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========================
# 1. Terima argument dari NestJS
# ========================
if len(sys.argv) < 3:
    print(f"Usage: {sys.argv[0]} <input_zip_path> <output_dir>", file=sys.stderr)
    sys.exit(1)

input_zip = sys.argv[1]      # Path ke ZIP dari NestJS
output_dir = sys.argv[2]     # Output directory dari NestJS
extract_dir = os.path.join(output_dir, "extracted")

print(f"[INFO] Input ZIP: {input_zip}")
print(f"[INFO] Output Dir: {output_dir}")
print(f"[INFO] Extract Dir: {extract_dir}")

# ========================
# 2. Validasi input file
# ========================
if not os.path.exists(input_zip):
    print(f"ERROR: Input file tidak ditemukan: {input_zip}", file=sys.stderr)
    sys.exit(1)

if not os.path.isfile(input_zip):
    print(f"ERROR: {input_zip} bukan file", file=sys.stderr)
    sys.exit(1)

# ========================
# 3. Create directories
# ========================
os.makedirs(extract_dir, exist_ok=True)
os.makedirs(output_dir, exist_ok=True)

try:
    # ========================
    # 4. Extract ZIP
    # ========================
    print(f"[INFO] Extracting ZIP file...")
    with zipfile.ZipFile(input_zip, 'r') as zip_ref:
        zip_ref.extractall(extract_dir)
    print(f"[INFO] ZIP extracted successfully")

    # ========================
    # 5. Ekstensi file gambar
    # ========================
    ekstensi_valid = [".jpg", ".jpeg", ".png", ".bmp"]

    # Ambil semua gambar dari folder extracted
    gambar_list = []
    for root, dirs, files in os.walk(extract_dir):
        for f in files:
            if os.path.splitext(f)[1].lower() in ekstensi_valid:
                gambar_list.append(os.path.join(root, f))
    gambar_list.sort()

    print(f"[INFO] Found {len(gambar_list)} images")
    if len(gambar_list) == 0:
        print("WARNING: No images found in ZIP", file=sys.stderr)

    for img_path in gambar_list:
        print(f"[INFO] Image: {img_path}")

    # ========================
    # 6. Buat dokumen Word baru
    # ========================
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(13)   # Legal/F4
    section.page_width = Inches(8.5)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # ========================
    # 7. Fungsi hapus border tabel
    # ========================
    def hapus_border_tabel(table):
        tbl = table._element
        tblPr = tbl.xpath(".//w:tblPr")
        if not tblPr:
            return
        tblPr = tblPr[0]
        tblBorders = tblPr.xpath(".//w:tblBorders")
        if tblBorders:
            tblPr.remove(tblBorders[0])
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    # ========================
    # 8. Loop tiap 4 gambar per halaman
    # ========================
    for i in range(0, len(gambar_list), 4):
        batch = gambar_list[i:i+4]

        # Tabel 2x2
        table = doc.add_table(rows=2, cols=2)
        table.autofit = False
        hapus_border_tabel(table)

        for j, gambar_path in enumerate(batch):
            row = j // 2
            col = j % 2
            cell = table.cell(row, col)
            paragraf = cell.paragraphs[0]
            paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraf.add_run()

            try:
                print(f"[INFO] Processing image: {gambar_path}")
                img = Image.open(gambar_path)

                # Rotasi otomatis landscape
                if img.width > img.height:
                    print(f"[INFO] Rotating image (landscape detected)")
                    img = img.rotate(90, expand=True)

                tmp_path = os.path.join(output_dir, f"tmp_rotated_{i}_{j}.png")
                img.save(tmp_path)

                run.add_picture(tmp_path, width=Inches(2.93), height=Inches(5.69))
                print(f"[INFO] Image added successfully")

            except Exception as e:
                print(f"ERROR: Gagal menyisipkan {gambar_path}: {e}", file=sys.stderr)

        if i + 4 < len(gambar_list):
            doc.add_page_break()

    # ========================
    # 9. Simpan hasil dokumen
    # ========================
    output_file = os.path.join(output_dir, "output.docx")
    print(f"[INFO] Saving document to: {output_file}")
    doc.save(output_file)

    # Verify file exists
    if os.path.exists(output_file):
        file_size = os.path.getsize(output_file)
        print(f"[SUCCESS] Document created: {output_file} (Size: {file_size} bytes)")
        sys.exit(0)
    else:
        print(f"ERROR: Failed to save document", file=sys.stderr)
        sys.exit(1)

except zipfile.BadZipFile as e:
    print(f"ERROR: Invalid ZIP file: {e}", file=sys.stderr)
    sys.exit(1)

except Exception as e:
    print(f"ERROR: {type(e).__name__}: {e}", file=sys.stderr)
    import traceback
    traceback.print_exc()
    sys.exit(1)